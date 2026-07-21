"""
Analizador de documentos de clientes - Motor principal
Soporta: PDF, JPG/PNG, DOCX, imágenes escaneadas (OCR)
Base de datos: PostgreSQL
"""

import os
import sys
import json
import base64
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional

import google.generativeai as genai
import psycopg2
from psycopg2.extras import RealDictCursor
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from pypdf import PdfReader
import pdfplumber
from docx import Document as DocxDocument
import io
from dotenv import load_dotenv

load_dotenv()

# ─── Configurar Tesseract (Windows) ───────────────────────────────────────────
import sys as _sys
if _sys.platform == "win32":
    import pytesseract as _tess
    _tess_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    import os as _os
    if _os.path.exists(_tess_path):
        _tess.pytesseract.tesseract_cmd = _tess_path

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ──────────────────────────────────────────────
# CONFIG
# ──────────────────────────────────────────────
DB_CONFIG = {
    "host": os.getenv("DB_HOST", "localhost"),
    "port": int(os.getenv("DB_PORT", 5432)),
    "dbname": os.getenv("DB_NAME", "postgres"),
    "user": os.getenv("DB_USER", "postgres"),
    "password": os.getenv("DB_PASSWORD", ""),
}

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".jpg": "imagen",
    ".jpeg": "imagen",
    ".png": "imagen",
    ".docx": "word",
    ".doc": "word",
}

DOC_TYPE_KEYWORDS = {
    "cedula": ["cedul", "cc ", "c.c", "identificacion", "id ", "dni"],
    "datacredito": ["datacredit", "data_credit", "buro", "bureau"],
    "cifin": ["cifin", "asobancaria"],
    "informe": ["confiab", "informe", "reporte", "analisis", "riesgo"],
}


# ──────────────────────────────────────────────
# BASE DE DATOS
# ──────────────────────────────────────────────
def get_connection():
    return psycopg2.connect(**DB_CONFIG)


def init_db():
    """Crea las tablas si no existen."""
    sql = """
    CREATE TABLE IF NOT EXISTS clientes_credito (
        id                      SERIAL PRIMARY KEY,
        nombre                  TEXT,
        apellidos               TEXT,
        nombres                 TEXT,
        cedula                  TEXT UNIQUE,
        fecha_nacimiento        TEXT,
        lugar_nacimiento        TEXT,
        fecha_expedicion        TEXT,
        lugar_expedicion        TEXT,
        sexo                    TEXT,
        estatura                TEXT,
        grupo_sanguineo         TEXT,
        score_datacredito       INTEGER,
        endeudamiento_datacredito NUMERIC(5,2),
        score_cifin             INTEGER,
        obligaciones_cifin      INTEGER,
        es_confiable            BOOLEAN,
        nivel_riesgo            TEXT CHECK (nivel_riesgo IN ('bajo','medio','alto')),
        resumen_ia              TEXT,
        documentos_fuente       TEXT,
        fecha_analisis          TIMESTAMP DEFAULT NOW(),
        raw_json                JSONB
    );

    CREATE TABLE IF NOT EXISTS documentos_procesados (
        id              SERIAL PRIMARY KEY,
        cliente_id      INTEGER REFERENCES clientes_credito(id) ON DELETE CASCADE,
        nombre_archivo  TEXT,
        tipo_documento  TEXT,
        formato         TEXT,
        texto_extraido  TEXT,
        fecha_proceso   TIMESTAMP DEFAULT NOW()
    );
    """
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(sql)
        conn.commit()
    log.info("Base de datos inicializada correctamente.")


def save_cliente(data: dict, doc_names: list[str], docs_info: list[dict]) -> int:
    """Inserta o actualiza un cliente en la base de datos."""
    sql_cliente = """
    INSERT INTO clientes_credito
        (nombre, apellidos, nombres, cedula, 
         fecha_nacimiento, lugar_nacimiento,
         fecha_expedicion, lugar_expedicion,
         sexo, estatura, grupo_sanguineo,
         score_datacredito, endeudamiento_datacredito,
         score_cifin, obligaciones_cifin,
         es_confiable, nivel_riesgo, resumen_ia,
         documentos_fuente, raw_json)
    VALUES
        (%(nombre)s, %(apellidos)s, %(nombres)s, %(cedula)s,
         %(fecha_nacimiento)s, %(lugar_nacimiento)s,
         %(fecha_expedicion)s, %(lugar_expedicion)s,
         %(sexo)s, %(estatura)s, %(grupo_sanguineo)s,
         %(score_datacredito)s, %(endeudamiento_datacredito)s,
         %(score_cifin)s, %(obligaciones_cifin)s,
         %(es_confiable)s, %(nivel_riesgo)s, %(resumen_ia)s,
         %(documentos_fuente)s, %(raw_json)s)
    ON CONFLICT (cedula) DO UPDATE SET
        nombre                    = EXCLUDED.nombre,
        apellidos                 = EXCLUDED.apellidos,
        nombres                   = EXCLUDED.nombres,
        fecha_nacimiento          = EXCLUDED.fecha_nacimiento,
        lugar_nacimiento          = EXCLUDED.lugar_nacimiento,
        fecha_expedicion          = EXCLUDED.fecha_expedicion,
        lugar_expedicion          = EXCLUDED.lugar_expedicion,
        sexo                      = EXCLUDED.sexo,
        estatura                  = EXCLUDED.estatura,
        grupo_sanguineo           = EXCLUDED.grupo_sanguineo,
        score_datacredito         = EXCLUDED.score_datacredito,
        endeudamiento_datacredito = EXCLUDED.endeudamiento_datacredito,
        score_cifin               = EXCLUDED.score_cifin,
        obligaciones_cifin        = EXCLUDED.obligaciones_cifin,
        es_confiable              = EXCLUDED.es_confiable,
        nivel_riesgo              = EXCLUDED.nivel_riesgo,
        resumen_ia                = EXCLUDED.resumen_ia,
        documentos_fuente         = EXCLUDED.documentos_fuente,
        fecha_analisis            = NOW(),
        raw_json                  = EXCLUDED.raw_json
    RETURNING id;
    """
    params = {
        "nombre": data.get("nombre"),
        "apellidos": data.get("apellidos"),
        "nombres": data.get("nombres"),
        "cedula": str(data.get("cedula", "")),
        "fecha_nacimiento": data.get("fecha_nacimiento"),
        "lugar_nacimiento": data.get("lugar_nacimiento"),
        "fecha_expedicion": data.get("fecha_expedicion"),
        "lugar_expedicion": data.get("lugar_expedicion"),
        "sexo": data.get("sexo"),
        "estatura": data.get("estatura"),
        "grupo_sanguineo": data.get("grupo_sanguineo"),
        "score_datacredito": data.get("score_datacredito"),
        "endeudamiento_datacredito": data.get("endeudamiento_datacredito"),
        "score_cifin": data.get("score_cifin"),
        "obligaciones_cifin": data.get("obligaciones_cifin"),
        "es_confiable": data.get("es_confiable"),
        "nivel_riesgo": data.get("nivel_riesgo", "medio"),
        "resumen_ia": data.get("resumen_ia"),
        "documentos_fuente": ", ".join(doc_names),
        "raw_json": json.dumps(data, ensure_ascii=False),
    }

    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute(sql_cliente, params)
            cliente_id = cur.fetchone()[0]

            # Guardar docs individuales
            for doc in docs_info:
                cur.execute(
                    """
                    INSERT INTO documentos_procesados
                        (cliente_id, nombre_archivo, tipo_documento, formato, texto_extraido)
                    VALUES (%s, %s, %s, %s, %s)
                    """,
                    (cliente_id, doc["nombre"], doc["tipo"], doc["formato"], doc.get("texto", "")[:5000]),
                )
        conn.commit()
    return cliente_id


# ──────────────────────────────────────────────
# EXTRACCIÓN DE TEXTO / IMÁGENES
# ──────────────────────────────────────────────
def detect_doc_type(filename: str) -> str:
    lower = filename.lower()
    for tipo, keywords in DOC_TYPE_KEYWORDS.items():
        if any(k in lower for k in keywords):
            return tipo
    return "otro"


import pypdfium2 as pdfium
import cv2
import numpy as np

from pdf417decoder import PDF417Decoder

def extract_pdf417(image_pil) -> list:
    """Detecta y decodifica códigos PDF417 (comunes en cédulas colombianas)."""
    try:
        # PDF417Decoder necesita un numpy array en formato BGR (OpenCV)
        img_rgb = np.array(image_pil.convert('RGB'))
        img_bgr = img_rgb[:, :, ::-1]  # RGB → BGR
        decoder = PDF417Decoder(img_bgr)
        count = decoder.decode()
        if count > 0:
            results = []
            for i in range(count):
                raw = decoder.barcode_data_index(i)
                results.append(raw if isinstance(raw, str) else raw.decode('utf-8', errors='ignore'))
            return results
    except Exception as e:
        log.debug(f"Fallo en decodificación PDF417: {e}")
    return []

def extract_qr(image_pil) -> list[str]:
    """Detecta y decodifica QRs con preprocesamiento para mejorar la detección."""
    # Convertir PIL a formato OpenCV (Grises)
    img = np.array(image_pil.convert('RGB'))
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    
    detector = cv2.QRCodeDetector()
    
    # 1. Intento normal
    retval, decoded_info, points, straight_qrcode = detector.detectAndDecodeMulti(gray)
    if retval and any(decoded_info):
        return [info for info in decoded_info if info]
    
    # 2. Intento con Umbralizado (para QR en fondos ruidosos)
    _, thresh = cv2.threshold(gray, 128, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    retval, decoded_info, points, straight_qrcode = detector.detectAndDecodeMulti(thresh)
    if retval and any(decoded_info):
        return [info for info in decoded_info if info]
        
    return []

def extract_from_pdf(path: Path) -> tuple[str, list]:
    """Extrae texto de PDF normal. Si falla, hace OCR. Usa pypdfium2 para imágenes."""
    text = ""
    images_b64 = []
    qr_data = []

    # Intento 1: texto digital
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                text += t + "\n"
    except Exception:
        pass
    
    if len(text.strip()) < 10:
        try:
            pdf = pdfium.PdfDocument(str(path))
            for page in pdf:
                text += page.get_textpage().get_text_range() + "\n"
            pdf.close()
        except Exception:
            pass

    # Intento 2: Convertir a imágenes para OCR y QR (300 DPI para mejor detalle)
    log.info(f"Convirtiendo PDF a imágenes (300 DPI): {path.name}")
    try:
        pdf = pdfium.PdfDocument(str(path))
        for i in range(len(pdf)):
            page = pdf[i]
            # Renderizar a 200 DPI
            bitmap = page.render(scale=200/72)
            img = bitmap.to_pil()
            
            # 1. Buscar QR y PDF417
            qrs = extract_qr(img)
            barcodes = extract_pdf417(img)
            
            if qrs:
                qr_data.extend(qrs)
                log.info(f"  [QR] Detectado en pág {i+1}")
            if barcodes:
                qr_data.append(str(barcodes))
                log.info(f"  [PDF417] Detectado en pág {i+1}")
            
            # 2. Si el texto digital fue pobre, aplicar OCR con preprocesamiento
            if len(text.strip()) < 50:
                text += ocr_image(img) + "\n"
            
            # 3. Guardar base64 para la IA
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            images_b64.append(base64.b64encode(buf.getvalue()).decode())
            
        pdf.close()
    except Exception as e:
        log.warning(f"Error procesando visualmente el PDF: {e}")

    if qr_data:
        text += "\n--- DATOS QR DETECTADOS ---\n" + "\n".join(qr_data)

    log.info(f"Análisis visual terminado. Texto extraído (largo={len(text)}): {text[:200]}...")
    return text, images_b64


def preprocess_for_ocr(pil_img) -> list:
    """
    Aplica múltiples estrategias de preprocesamiento y devuelve todas las
    versiones para que el OCR tenga más chances de leer el texto.
    """
    gray = np.array(pil_img.convert('L'))
    versions = [gray]

    # Filtro bilateral (suaviza ruido preservando bordes de letras)
    bilat = cv2.bilateralFilter(gray, 11, 17, 17)
    versions.append(bilat)

    # CLAHE (mejora contraste adaptativo en zonas oscuras)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(bilat)
    versions.append(enhanced)

    # Sharpen
    kernel = np.array([[-1, -1, -1], [-1, 9, -1], [-1, -1, -1]])
    sharp = cv2.filter2D(enhanced, -1, kernel)
    versions.append(sharp)

    return [Image.fromarray(v) for v in versions]


def ocr_image(pil_img) -> str:
    """OCR con múltiples estrategias de preprocesamiento. Devuelve el mejor resultado."""
    if not pytesseract:
        return ""
    config = '--oem 3 --psm 6 -l spa'
    best_text = ""
    best_score = 0

    try:
        for version in preprocess_for_ocr(pil_img):
            text = pytesseract.image_to_string(version, config=config)
            score = sum(1 for c in text if c.isalpha())
            if score > best_score:
                best_score = score
                best_text = text
    except Exception as e:
        log.warning(f"  [OCR] Falló: {e}")

    return best_text


def extract_from_image(path: Path, skip_ocr: bool = False) -> tuple[str, str]:
    """Extrae texto via OCR con preprocesamiento, busca QRs y devuelve base64."""
    img = Image.open(path)
    text = ""
    if not skip_ocr:
        text = ocr_image(img)
    
    # Buscar QR y PDF417
    qrs = extract_qr(img)
    barcodes = extract_pdf417(img)
    
    if qrs or barcodes:
        qr_text = "\n".join(qrs) + ("\n" + str(barcodes) if barcodes else "")
        text += "\n--- DATOS QR DETECTADOS ---\n" + qr_text
        log.info(f"  [Barcodes] Detectados en imagen {path.name}")

    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="JPEG", quality=85)
    b64 = base64.b64encode(buf.getvalue()).decode()
    return text, b64



def extract_from_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ──────────────────────────────────────────────
# ANÁLISIS CON CLAUDE
# ──────────────────────────────────────────────
SYSTEM_PROMPT = """Eres un experto en análisis de documentos financieros colombianos.
Recibirás documentos de un cliente (cédula, consulta DataCrédito, CIFIN, informes de confiabilidad).
Tu tarea es extraer y consolidar toda la información relevante.
Responde ÚNICAMENTE con un objeto JSON válido, sin texto adicional, sin bloques markdown."""

JSON_SCHEMA = """{
  "apellidos": "apellidos del cliente",
  "nombres": "nombres del cliente",
  "cedula": "número de cédula sin puntos",
  "fecha_nacimiento": "YYYY-MM-DD o null",
  "lugar_nacimiento": "Ciudad, Departamento o null",
  "fecha_expedicion": "YYYY-MM-DD o null",
  "lugar_expedicion": "Ciudad, Departamento o null",
  "sexo": "M o F o null",
  "estatura": "ej: 1.70 o null",
  "grupo_sanguineo": "ej: O+ o null",
  "score_datacredito": número entero o null,
  "endeudamiento_datacredito": porcentaje como número decimal o null,
  "score_cifin": número entero o null,
  "obligaciones_cifin": número entero de obligaciones o null,
  "es_confiable": true o false,
  "nivel_riesgo": "bajo" | "medio" | "alto",
  "resumen_ia": "resumen de 2-3 oraciones del perfil crediticio"
}"""


import re

def parse_barcode_data(raw_text: str) -> dict:
    """
    Intenta parsear los datos de una cédula colombiana desde el texto del PDF417/QR.
    El formato suele ser binario o con separadores nulos.
    """
    data = {}
    try:
        # Limpieza básica
        clean = "".join(ch if ch.isalnum() or ch in " -|_" else " " for ch in raw_text)
        
        # 1. Buscar Cédula (secuencia de 8 a 10 dígitos)
        id_match = re.search(r'\b(\d{8,10})\b', clean)
        if id_match:
            data["cedula"] = id_match.group(1)
            
        # 2. Buscar Nombres y Apellidos
        # En el formato PDF417 de la Registraduría, los nombres vienen en mayúsculas
        # después del número de cédula, separados por espacios o caracteres nulos.
        parts = [p for p in re.split(r'[^A-ZÁÉÍÓÚÑ]', clean) if len(p) > 1]
        if len(parts) >= 2:
            # Típicamente: APELLIDO1 APELLIDO2 NOMBRE1 NOMBRE2
            # Intentamos reconstruir el nombre completo
            nombre_completo = " ".join(parts[:4])
            data["nombre"] = nombre_completo
            
        # 3. Fecha de nacimiento (opcional, si viene en formato YYYYMMDD)
        date_match = re.search(r'\b(19|20)\d{6}\b', clean)
        if date_match:
            # Podríamos parsearla pero por ahora la IA o el manual la manejarán
            pass
            
    except Exception as e:
        log.warning(f"Error parseando datos del código de barras: {e}")
    return data

MESES_ES = {"ENE": "01", "FEB": "02", "MAR": "03", "ABR": "04", "MAY": "05", "JUN": "06",
            "JUL": "07", "AGO": "08", "SEP": "09", "OCT": "10", "NOV": "11", "DIC": "12"}

def parse_date_es(text: str) -> str | None:
    """Convierte fechas colombianas como '29-JUL-1988' a formato YYYY-MM-DD."""
    m = re.search(r'(\d{1,2})[-/]([A-Z]{3})[-/](\d{4})', text.upper())
    if m:
        d, mes, y = m.group(1), m.group(2), m.group(3)
        mm = MESES_ES.get(mes)
        if mm:
            return f"{y}-{mm}-{d.zfill(2)}"
    # Fecha numérica normal
    m2 = re.search(r'(\d{4})-(\d{2})-(\d{2})', text)
    if m2:
        return m2.group(0)
    return None

def extract_data_manually(text: str, filename: str = "") -> dict:
    """
    Extrae información usando Regex tolerantes al ruido OCR.
    Prioriza coincidencias con contexto (etiquetas) y luego heurísticas.
    """
    data = {
        # Campos individuales de la cédula colombiana
        "apellidos": None,
        "nombres": None,
        "cedula": None,
        "fecha_nacimiento": None,
        "lugar_nacimiento": None,
        "fecha_expedicion": None,
        "lugar_expedicion": None,
        "sexo": None,
        "estatura": None,
        "grupo_sanguineo": None,
        # Campos financieros (se llenan con otros documentos)
        "score_datacredito": 0,
        "endeudamiento_datacredito": 0,
        "score_cifin": 0,
        "es_confiable": True,
        "nivel_riesgo": "bajo",
        "resumen_ia": "Extraído mediante OCR y reglas locales (Sin IA)."
    }

    # ── 0. Limpieza de Ruido MRZ / Parte Inferior ─────────────────────────────
    # El bloque MRZ suele empezar con I<COL o similar y contiene mucha basura OCR
    text = re.sub(r'(?m)^[I|1|L|7]<[A-Z0-9<]{10,}.*$', '', text)
    # También eliminar líneas que tengan demasiados caracteres '<' seguidos
    text = '\n'.join([line for line in text.split('\n') if line.count('<') < 4])

    # ── 1. Cédula ──────────────────────────────────────────────────────────────
    id_m = re.search(r'(?:C[EÉ]DULA|CC|N[ÚU]MERO)[:.\s]*([\d.,\s]{7,15})', text, re.IGNORECASE)
    if id_m:
        data["cedula"] = re.sub(r'[^0-9]', '', id_m.group(1))
    else:
        nums = re.findall(r'\b(\d[\d.]{6,11}\d)\b', text)
        for n in nums:
            clean = re.sub(r'[^0-9]', '', n)
            if 7 <= len(clean) <= 10:
                data["cedula"] = clean
                break
    if not data["cedula"] and filename:
        fn_m = re.search(r'(\d{7,10})', filename)
        if fn_m:
            data["cedula"] = fn_m.group(1)
            log.info(f"  [Fallback] Cédula del nombre de archivo: {data['cedula']}")

    # ── 2. Apellidos ───────────────────────────────────────────────────────────
    # Buscar etiqueta APELLIDOS y capturar 1 o 2 líneas
    ap_m = re.search(r'APELLIDOS?[:.\s]+([A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,})?)', text, re.IGNORECASE)
    if ap_m:
        data["apellidos"] = ap_m.group(1).strip()

    # ── 3. Nombres ─────────────────────────────────────────────────────────────
    # Buscar etiqueta NOMBRES y capturar 1 o 2 líneas
    no_m = re.search(r'NOMBRES?[:.\s]+([A-ZÁÉÍÓÚÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,})?)', text, re.IGNORECASE)
    if no_m:
        data["nombres"] = no_m.group(1).strip()

    # Si no hubo etiquetas, buscar bloque de palabras en mayúsculas
    if not data["apellidos"] and not data["nombres"]:
        STOPWORDS = {"REPUBLICA", "COLOMBIA", "CEDULA", "CIUDADANIA", "IDENTIFICACION",
                     "PERSONAL", "FECHA", "LUGAR", "NACIMIENTO", "EXPEDICION",
                     "SEXO", "ESTATURA", "REGISTRADOR", "NACIONAL", "ATLANTICO"}
        words = re.findall(r'\b([A-ZÁÉÍÓÚÑ]{3,})\b', text)
        name_words = [w for w in words if w not in STOPWORDS]
        if len(name_words) >= 4:
            data["apellidos"] = " ".join(name_words[:2])
            data["nombres"]   = " ".join(name_words[2:4])
        elif len(name_words) >= 2:
            data["apellidos"] = " ".join(name_words[:2])

    # ── 4. Todas las fechas (nacimiento + expedición) ──────────────────────────
    all_dates = []
    for m in re.finditer(r'\d{1,2}[-/][A-Z]{3}[-/]\d{4}', text.upper()):
        d = parse_date_es(m.group(0))
        if d:
            all_dates.append((m.start(), d))
    # Cédula colombiana: 1ª fecha = nacimiento, 2ª = expedición
    if len(all_dates) >= 2:
        data["fecha_nacimiento"] = all_dates[0][1]
        data["fecha_expedicion"] = all_dates[1][1]
    elif len(all_dates) == 1:
        data["fecha_expedicion"] = all_dates[0][1]

    # ── 5. Lugar de nacimiento ─────────────────────────────────────────────────
    nac_m = re.search(
        r'(?:LUGAR\s+DE\s+NACIMIENTO|LUGAR\s+NAC)[:.\s]+([A-ZÁÉÍÓÚÑ\s]{3,30})',
        text, re.IGNORECASE
    )
    if nac_m:
        data["lugar_nacimiento"] = nac_m.group(1).strip()

    # ── 6. Lugar de expedición ─────────────────────────────────────────────────
    exp_m = re.search(
        r'(?:EXPEDICI[OÓ]N|FECHA\s+Y\s+LUGAR)[:.\s]+(?:\d{1,2}[-/][A-Z]{3}[-/]\d{4}\s+)?([A-ZÁÉÍÓÚÑ]{3,})',
        text, re.IGNORECASE
    )
    if exp_m:
        data["lugar_expedicion"] = exp_m.group(1).strip()
    else:
        CIUDADES = ["BARRANQUILLA", "BOGOTA", "BOGOTÁ", "MEDELLIN", "MEDELLÍN",
                    "CALI", "CARTAGENA", "CUCUTA", "BUCARAMANGA", "PEREIRA",
                    "MANIZALES", "SANTA MARTA", "IBAGUE", "VILLAVICENCIO"]
        for ciudad in CIUDADES:
            if ciudad in text.upper():
                data["lugar_expedicion"] = ciudad.title()
                if not data["lugar_nacimiento"]:
                    data["lugar_nacimiento"] = ciudad.title()
                break

    # ── 7. Sexo ────────────────────────────────────────────────────────────────
    # En la parte de atrás dice "SEXO M" o "SEXO F"
    sex_m = re.search(r'SEXO[:.\s]*([MF])\b', text, re.IGNORECASE)
    if sex_m:
        data["sexo"] = sex_m.group(1).upper()
    elif not data["sexo"]:
        # Búsqueda libre de M o F aislados cerca de etiquetas de estatura
        if re.search(r'(?:SEXO|ESTATURA).{1,10}\bM\b', text, re.IGNORECASE): data["sexo"] = "M"
        elif re.search(r'(?:SEXO|ESTATURA).{1,10}\bF\b', text, re.IGNORECASE): data["sexo"] = "F"

    # ── 8. Estatura ────────────────────────────────────────────────────────────
    # Ej: 1.75 o 1,75
    est_m = re.search(r'(?:ESTATURA|EST)[:.\s]*([12][.,]\d{2})\b', text, re.IGNORECASE)
    if est_m:
        data["estatura"] = est_m.group(1).replace(',', '.')

    # ── 9. Grupo Sanguíneo ─────────────────────────────────────────────────────
    # Ej: O+, A-, AB+
    gs_m = re.search(r'G\.\s?S[:.\s]*([ABO]{1,2}[+-])', text, re.IGNORECASE)
    if not gs_m:
        gs_m = re.search(r'\b([ABO]{1,2}[+-]|O[+-])\b', text)
    if gs_m:
        data["grupo_sanguineo"] = gs_m.group(1).replace(' ', '').upper()

    # ── 10. Datos financieros (DataCrédito / CIFIN) ───────────────────────────
    # Score DataCrédito (3 dígitos, típicamente 150-950)
    dc_m = re.search(r'(?:PUNTAJE|SCORE|PUNTUACI[OÓ]N)[.\s]*(?:DATACR[EÉ]DITO)?[:.\s]*(\d{3})\b', text, re.IGNORECASE)
    if dc_m:
        data["score_datacredito"] = int(dc_m.group(1))
    
    # Endeudamiento (Porcentaje)
    end_m = re.search(r'(?:ENDEUDAMIENTO|CAPACIDAD)[.\s]*(?:TOTAL|PAGO)?[:.\s]*(\d{1,3}(?:[.,]\d{1,2})?)\s*%', text, re.IGNORECASE)
    if end_m:
        data["endeudamiento_datacredito"] = float(end_m.group(1).replace(',', '.'))

    # Cédula en el reporte financiero (opcional para auto-asociar)
    # Ej: "Identificación: 12.345.678"
    id_fin_m = re.search(r'(?:IDENTIFICACI[OÓ]N|C[EÉ]DULA|DOCUMENTO)[:.\s]*([\d.,\s]{7,15})', text, re.IGNORECASE)
    if id_fin_m:
        val = re.sub(r'[^0-9]', '', id_fin_m.group(1))
        if 7 <= len(val) <= 10:
            data["cedula"] = val

    # ── 11. Campos Adicionales Buró (Nuevos) ──────────────────────────────────
    # Estado Cédula (Vigente/Cancelada)
    if re.search(r'CEDULA\s+VIGENTE', text, re.IGNORECASE): data["estado_cedula"] = "VIGENTE"
    
    # Moras
    mora_m = re.search(r'(?:MORA\s+M[AÁ]XIMA|M[AÁ]XIMO\s+ATRASO|HISTORIA\s+DE\s+MORA)[:.\s]*(\d+)', text, re.IGNORECASE)
    if mora_m: data["mora_maxima"] = int(mora_m.group(1))
    
    # Huellas de consulta (últimos 6 meses)
    huellas_m = re.search(r'(\d+)\s*(?:HUELLAS|CONSULTAS|CONSULTADA)\s+EN\s+LOS\s+[UÚ]LTIMOS', text, re.IGNORECASE)
    if huellas_m: data["huellas_consulta"] = int(huellas_m.group(1))
    
    # Saldo Total y Cupo Total (Limpieza de moneda)
    def clean_currency(val_str):
        return int(re.sub(r'[^0-9]', '', val_str))

    saldo_m = re.search(r'SALDO\s+TOTAL[:.\s]*\$?\s*([\d.,\s]{5,20})', text, re.IGNORECASE)
    if saldo_m: data["saldo_total"] = clean_currency(saldo_m.group(1))
    
    cupo_m = re.search(r'CUPO\s+TOTAL[:.\s]*\$?\s*([\d.,\s]{5,20})', text, re.IGNORECASE)
    if cupo_m: data["cupo_total"] = clean_currency(cupo_m.group(1))

    # Cuentas Abiertas
    ab_m = re.search(r'(\d+)\s+CUENTAS\s+ABIERTAS', text, re.IGNORECASE)
    if ab_m: data["cuentas_abiertas"] = int(ab_m.group(1))
    end_m = re.search(r'(?:ENDEUDAMIENTO|CAPACIDAD\s+DE\s+PAGO)[:.\s]*(\d{1,2}[.,]\d{1,2})\s*%', text, re.IGNORECASE)
    if end_m:
        data["endeudamiento_datacredito"] = float(end_m.group(1).replace(',', '.'))

    # Score CIFIN
    cf_m = re.search(r'(?:SCORE|PUNTAJE)\s+CIFIN[:.\s]*(\d{3})\b', text, re.IGNORECASE)
    if cf_m:
        data["score_cifin"] = int(cf_m.group(1))

    # Obligaciones CIFIN (ej: "5 obligaciones")
    obl_m = re.search(r'(\d{1,2})\s+OBLIGACIONES\b', text, re.IGNORECASE)
    if obl_m:
        data["obligaciones_cifin"] = int(obl_m.group(1))
    else:
        # Buscar en tablas o listas
        obl_m2 = re.search(r'TOTAL\s+OBLIGACIONES[:.\s]*(\d{1,2})\b', text, re.IGNORECASE)
        if obl_m2:
            data["obligaciones_cifin"] = int(obl_m2.group(1))

    # ── 11. Limpieza de Ruido (Opcional) ──────────────────────────────────────
    # El usuario solicita descartar información de la parte de atrás inferior
    # que suele ser el bloque MRZ (I<COL...) o etiquetas de fondo.
    # No eliminamos el texto del todo para no romper regex, pero podemos ignorar 
    # coincidencias si el texto parece basura técnica.
    
    # ── 12. Compatibilidad: campo "nombre" como apellidos + nombres ───────────
    data["nombre"] = " ".join(filter(None, [data.get("apellidos"), data.get("nombres")])) or None

    log.info(f"Datos extraídos: '{data['apellidos']}' '{data['nombres']}' CC={data['cedula']} DC={data['score_datacredito']}")
    return data


def analyze_with_claude(docs_info: list[dict]) -> dict:
    """Envía los documentos a Gemini para un análisis profundo. Fallback a manual si no hay API KEY."""
    api_key = os.getenv("GEMINI_API_KEY", "") or os.getenv("ANTHROPIC_API_KEY", "")
    
    if not api_key or "PONER_AQUI" in api_key:
        log.warning("No hay API KEY de Gemini. Usando extracción por reglas locales.")
        full_text = "\n".join(d.get("texto", "") for d in docs_info)
        # Usar el nombre del primer archivo como referencia
        filename = docs_info[0].get("nombre", "") if docs_info else ""
        return extract_data_manually(full_text, filename)

    genai.configure(api_key=api_key)

    content = []

    for doc in docs_info:
        tipo_label = f"[{doc['tipo'].upper()}] Archivo: {doc['nombre']}"

        # Agregar imágenes si hay
        for b64 in doc.get("images_b64", []):
            content.append({
                "mime_type": "image/jpeg",
                "data": base64.b64decode(b64)
            })

        # Agregar texto extraído
        if doc.get("texto"):
            content.append(f"{tipo_label}\nTexto extraído:\n{doc['texto'][:3000]}")
        else:
            content.append(tipo_label)

    content.append(f"""Analiza todos los documentos anteriores y extrae la información del cliente.
Responde ÚNICAMENTE con este JSON (sin markdown, sin explicaciones):
{JSON_SCHEMA}""")

    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        system_instruction=SYSTEM_PROMPT
    )

    response = model.generate_content(
        contents=content,
        generation_config={"response_mime_type": "application/json"}
    )

    raw = response.text.strip()
    raw = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)


def process_single_file(file_path: str) -> Optional[dict]:
    f = Path(file_path)
    if not f.exists():
        log.error(f"El archivo no existe: {file_path}")
        return None

    ext = f.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        log.error(f"Extensión no soportada: {ext}")
        return None

    fmt = SUPPORTED_EXTENSIONS[ext]
    tipo = detect_doc_type(f.name)
    log.info(f"Procesando archivo individual: {f.name} ({fmt}, tipo: {tipo})")

    doc = {"nombre": f.name, "tipo": tipo, "formato": fmt, "texto": "", "images_b64": []}

    if fmt == "pdf":
        text, imgs = extract_from_pdf(f)
        doc["texto"] = text
        doc["images_b64"] = imgs
    elif fmt == "imagen":
        # Si es cédula, la IA (Gemini) lee la imagen directamente. El OCR local no es necesario y es muy lento.
        text, b64 = extract_from_image(f, skip_ocr=(tipo == "cedula"))
        doc["texto"] = text
        doc["images_b64"] = [b64]
    elif fmt == "word":
        doc["texto"] = extract_from_docx(f)

    docs_info = [doc]

    log.info("Enviando documento a Claude para análisis...")
    try:
        data = analyze_with_claude(docs_info)
    except Exception as e:
        log.error(f"Error en análisis con Claude: {e}")
        return None

    log.info(f"Información extraída: {data.get('nombre')} (CC {data.get('cedula')})")
    
    # Nota: No guardamos automáticamente aquí si el usuario quiere confirmar los datos primero.
    # Pero el endpoint de la API puede decidir llamar a save_cliente después.
    return data





# ──────────────────────────────────────────────
# PROCESO PRINCIPAL
# ──────────────────────────────────────────────
def process_folder(folder_path: str) -> Optional[dict]:
    folder = Path(folder_path)
    if not folder.is_dir():
        log.error(f"La carpeta no existe: {folder_path}")
        return None

    files = [f for f in folder.iterdir() if f.suffix.lower() in SUPPORTED_EXTENSIONS]
    if not files:
        log.warning(f"No se encontraron documentos soportados en {folder_path}")
        return None

    log.info(f"Procesando {len(files)} archivo(s) en: {folder.name}")
    docs_info = []

    for f in files:
        ext = f.suffix.lower()
        fmt = SUPPORTED_EXTENSIONS[ext]
        tipo = detect_doc_type(f.name)
        log.info(f"  → {f.name} ({fmt}, tipo: {tipo})")

        doc = {"nombre": f.name, "tipo": tipo, "formato": fmt, "texto": "", "images_b64": []}

        if fmt == "pdf":
            text, imgs = extract_from_pdf(f)
            doc["texto"] = text
            doc["images_b64"] = imgs

        elif fmt == "imagen":
            text, b64 = extract_from_image(f)
            doc["texto"] = text
            doc["images_b64"] = [b64]

        elif fmt == "word":
            doc["texto"] = extract_from_docx(f)

        docs_info.append(doc)

    log.info("Enviando documentos a Claude para análisis...")
    try:
        data = analyze_with_claude(docs_info)
    except Exception as e:
        log.error(f"Error en análisis con Claude: {e}")
        return None

    log.info(f"Cliente identificado: {data.get('nombre')} (CC {data.get('cedula')})")

    cliente_id = save_cliente(data, [d["nombre"] for d in docs_info], docs_info)
    log.info(f"Guardado en PostgreSQL con ID: {cliente_id}")

    data["_cliente_id"] = cliente_id
    return data


# ──────────────────────────────────────────────
# ENTRY POINT
# ──────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python analyzer.py <carpeta_cliente>")
        print("Ejemplo: python analyzer.py ./clientes/carlos_perez")
        sys.exit(1)

    init_db()
    result = process_folder(sys.argv[1])

    if result:
        print("\n✅ Análisis completado:")
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("❌ No se pudo procesar la carpeta.")
        sys.exit(1)
